import csv
from project_path import *
import pandas as pd
import docx

# Query all the Order objects
orders = Order.objects.all()

# Define the CSV file path
csv_file_path = 'orders.csv'

# Write the data to a CSV file
with open(csv_file_path, 'w', newline='') as csv_file:
    csv_writer = csv.writer(csv_file)
    
    # Write the header row
    header = [
        'User',
        'Product Categories',
        'Product Sub Categories',
        'Product',
        'Price',
        'Quantity',
        'Total Price',
        'Date',
        'Slug',
    ]
    csv_writer.writerow(header)
    
    # Write the data rows
    for order in orders:
        csv_writer.writerow([
            order.user.username if order.user else '',
            order.product_categories.name if order.product_categories else '',
            order.product_sub_categories.name if order.product_sub_categories else '',
            order.product.name if order.product else '',
            order.price if order.price else '',
            str(order.quantity),
            str(order.total_price) if order.total_price else '',
            str(order.date) if order.date else '',
            order.slug if order.slug else '',
        ])

data = {
    'User': [order.user.username if order.user else '' for order in orders],
    'Product Categories': [order.product_categories.name if order.product_categories else '' for order in orders],
    'Product Sub Categories': [order.product_sub_categories.name if order.product_sub_categories else '' for order in orders],
    'Product': [order.product.name if order.product else '' for order in orders],
    'Price': [str(order.price) if order.price else '' for order in orders],
    'Quantity': [str(order.quantity) for order in orders],
    'Total Price': [str(order.total_price) if order.total_price else '' for order in orders],
    'Date': [str(order.date) if order.date else '' for order in orders],
    'Slug': [order.slug if order.slug else '' for order in orders],
}
df = pd.DataFrame(data)

# Define the Excel file path
excel_file_path = 'orders.xlsx'

# Write the DataFrame to an Excel file
df.to_excel(excel_file_path, index=False)

doc = docx.Document()

# Add a table to the document
table = doc.add_table(rows=1, cols=9)  # Adjust the number of columns as needed
def set_cell_style(cell, text, is_header=False):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.bold = is_header  # Make headers bold
    font = run.font
    font.name = 'Times New Roman'  # Set font to Times New Roman

# Define the table header row
header_cells = table.rows[0].cells
set_cell_style(header_cells[0], 'User', is_header=True)
set_cell_style(header_cells[1], 'Product Categories', is_header=True)
set_cell_style(header_cells[2], 'Product Sub Categories', is_header=True)
set_cell_style(header_cells[3], 'Product', is_header=True)
set_cell_style(header_cells[4], 'Price', is_header=True)
set_cell_style(header_cells[5], 'Quantity', is_header=True)
set_cell_style(header_cells[6], 'Total Price', is_header=True)
set_cell_style(header_cells[7], 'Date', is_header=True)
set_cell_style(header_cells[8], 'Slug', is_header=True)

# Fill in the table with data
for order in orders:
    row_cells = table.add_row().cells
    set_cell_style(row_cells[0], order.user.username if order.user else '')
    set_cell_style(row_cells[1], order.product_categories.name if order.product_categories else '')
    set_cell_style(row_cells[2], order.product_sub_categories.name if order.product_sub_categories else '')
    set_cell_style(row_cells[3], order.product.name if order.product else '')
    set_cell_style(row_cells[4], str(order.price) if order.price else '')
    set_cell_style(row_cells[5], str(order.quantity))
    set_cell_style(row_cells[6], str(order.total_price) if order.total_price else '')
    set_cell_style(row_cells[7], str(order.date) if order.date else '')
    set_cell_style(row_cells[8], order.slug if order.slug else '')

# Save the Word document
word_file_path = 'orders.docx'
doc.save(word_file_path)

print(f"Data from the 'Order' model has been exported to '{word_file_path}'.")
print(f"Data from the 'Order' model has been exported to '{excel_file_path}'.")
print(f"Data from the 'Order' model has been exported to '{csv_file_path}'.")
